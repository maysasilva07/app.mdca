import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import networkx as nx
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import tempfile
import os

# ------------------------------------------------------------
# CONFIGURAÇÃO DE FONTE PARA GARANTIR ACENTUAÇÃO
# ------------------------------------------------------------
plt.rcParams['font.family'] = 'DejaVu Sans'

# ------------------------------------------------------------
# MODELAGEM MATEMÁTICA: CRITIC E PROMETHEE
# ------------------------------------------------------------

def normalizar_critic(df, criterios_tipo):
    df_norm = pd.DataFrame(index=df.index)
    for col in df.columns:
        r_min = df[col].min()
        r_max = df[col].max()
        if r_max == r_min:
            df_norm[col] = 0.0
            continue
        if criterios_tipo[col] == 'max':
            df_norm[col] = (df[col] - r_min) / (r_max - r_min)
        else:
            df_norm[col] = (r_max - df[col]) / (r_max - r_min)
    return df_norm

def metodo_critic(df_norm):
    desvios = df_norm.std(ddof=1)
    correlacao = df_norm.corr()
    info_c = pd.Series(index=df_norm.columns, dtype=float)
    for col in df_norm.columns:
        soma_conflito = (1 - correlacao[col]).sum()
        info_c[col] = desvios[col] * soma_conflito
    pesos = info_c / info_c.sum()
    return pesos.to_dict(), desvios, correlacao, info_c

def funcao_preferencia_tipo_v(d, q, p):
    if d <= q:
        return 0.0
    elif d >= p:
        return 1.0
    else:
        return (d - q) / (p - q)

def promethee(df, pesos, criterios_tipo, q_limites, p_limites):
    n = len(df)
    alternativas = list(df.index)
    criterios = list(df.columns)

    pref_agregada = pd.DataFrame(0.0, index=alternativas, columns=alternativas)

    for a in alternativas:
        for b in alternativas:
            if a == b:
                continue
            soma_pi = 0.0
            for crit in criterios:
                if criterios_tipo[crit] == 'max':
                    d = df.loc[a, crit] - df.loc[b, crit]
                else:
                    d = df.loc[b, crit] - df.loc[a, crit]

                q = q_limites[crit]
                p = p_limites[crit]

                if p == q:
                    pref = 1.0 if d > q else 0.0
                else:
                    pref = funcao_preferencia_tipo_v(d, q, p)

                soma_pi += pesos[crit] * pref

            pref_agregada.loc[a, b] = soma_pi

    phi_mais = pref_agregada.sum(axis=1) / (n - 1)
    phi_menos = pref_agregada.sum(axis=0) / (n - 1)
    phi_liquido = phi_mais - phi_menos
    ranking = phi_liquido.sort_values(ascending=False).rank(ascending=False, method='min').astype(int)

    return pref_agregada, phi_mais, phi_menos, phi_liquido, ranking

# ------------------------------------------------------------
# FUNÇÕES DE VISUALIZAÇÃO
# ------------------------------------------------------------

def gerar_grafico_pesos(pesos):
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar(pesos.keys(), pesos.values(), color='royalblue')
    ax.set_title('Peso dos Critérios (CRITIC)')
    ax.set_ylabel('Peso (%)')
    plt.xticks(rotation=45)
    return fig

def gerar_grafico_fluxos(phi_mais, phi_menos, phi_liquido):
    df_plot = pd.DataFrame({'Fluxo+': phi_mais, 'Fluxo-': phi_menos, 'Fluxo Líquido': phi_liquido})
    fig, ax = plt.subplots(figsize=(10, 5))
    df_plot.plot(kind='bar', ax=ax, color=['green', 'red', 'blue'])
    ax.set_title('Fluxos de Superação - PROMETHEE')
    ax.set_ylabel('Valores dos Fluxos')
    plt.xticks(rotation=0)
    ax.axhline(0, color='black', linewidth=0.8)
    return fig

# ------------------------------------------------------------
# FUNÇÃO CORRIGIDA DO GRAFO (setas triangulares visíveis)
# ------------------------------------------------------------
def gerar_grafo_sobreclassificacao(phi_mais, phi_menos):
    """
    Gera o grafo de sobreclassificação (PROMETHEE I) com setas triangulares na ponta.
    A seta aponta de quem domina para quem é dominado.
    """
    G = nx.DiGraph()
    alt_list = list(phi_mais.index)
    G.add_nodes_from(alt_list)

    # Adiciona arestas baseado nos critérios do PROMETHEE I
    for a in alt_list:
        for b in alt_list:
            if a == b:
                continue
            phi_mais_a = phi_mais[a]
            phi_mais_b = phi_mais[b]
            phi_menos_a = phi_menos[a]
            phi_menos_b = phi_menos[b]

            cond1 = (phi_mais_a > phi_mais_b) and (phi_menos_a < phi_menos_b)
            cond2 = (phi_mais_a == phi_mais_b) and (phi_menos_a < phi_menos_b)
            cond3 = (phi_mais_a > phi_mais_b) and (phi_menos_a == phi_menos_b)

            if cond1 or cond2 or cond3:
                G.add_edge(a, b)

    fig, ax = plt.subplots(figsize=(14, 10))

    # Se não houver arestas, exibe um aviso no gráfico
    if G.number_of_edges() == 0:
        ax.text(0.5, 0.5, 'Nenhuma relação de sobreclassificação encontrada', 
                ha='center', va='center', fontsize=16, transform=ax.transAxes)
        ax.set_title("Grafo de Sobreclassificação (PROMETHEE I) - Sem relações", fontsize=16)
        ax.axis('off')
        return fig

    # Layout spring para distribuir bem os nós
    pos = nx.spring_layout(G, seed=42, k=3, iterations=100)

    # Desenha nós com bordas e tamanho ajustado
    nx.draw_networkx_nodes(G, pos, node_size=4000, node_color='lightblue', 
                           edgecolors='black', linewidths=2, ax=ax)

    # Desenha arestas com setas triangulares grandes e margens para garantir visibilidade
    nx.draw_networkx_edges(
        G, pos,
        arrows=True,
        arrowstyle='-|>',                 # Seta triangular preenchida
        arrowsize=40,                      # Tamanho da seta (bem visível)
        edge_color='gray',
        width=3,
        connectionstyle='arc3,rad=0.1',    # Curvatura leve para evitar sobreposição
        ax=ax,
        min_source_margin=25,               # Margem na origem (seta não começa dentro do nó)
        min_target_margin=25                 # Margem no destino (ponta não fica escondida)
    )

    # Desenha rótulos com fonte maior
    nx.draw_networkx_labels(G, pos, font_size=16, font_weight='bold', 
                            font_family='DejaVu Sans', ax=ax)

    ax.set_title("Grafo de Sobreclassificação (PROMETHEE I)", fontsize=18, pad=20)
    ax.axis('off')
    plt.tight_layout()
    return fig

# ------------------------------------------------------------
# FUNÇÕES DE RELATÓRIO (PDF E DOCX)
# ------------------------------------------------------------

def gerar_relatorio_pdf(dados, nome_arquivo):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Relatório MCDA - CRITIC + PROMETHEE", ln=True, align='C')
    pdf.ln(10)

    def ascii_only(texto):
        if isinstance(texto, str):
            return texto.encode('ascii', 'ignore').decode('ascii')
        return str(texto)

    def adicionar_tabela_texto(titulo, df):
        titulo_ascii = ascii_only(titulo)
        pdf.set_font("Arial", 'B', 10)
        pdf.cell(200, 10, txt=titulo_ascii, ln=True)
        pdf.set_font("Arial", size=8)
        linhas = ascii_only(df.to_string(index=True)).split('\n')
        for linha in linhas:
            if linha.strip():
                pdf.cell(200, 4, txt=linha, ln=True)
        pdf.ln(5)

    if 'q_limites' in dados and 'p_limites' in dados and 'tipos' in dados:
        pdf.set_font("Arial", 'B', 10)
        pdf.cell(200, 10, txt="Limiares dos Critérios (q e p):", ln=True)
        pdf.set_font("Arial", size=8)
        pdf.cell(40, 5, "Critério", border=1)
        pdf.cell(20, 5, "Tipo", border=1)
        pdf.cell(30, 5, "q (indif)", border=1)
        pdf.cell(30, 5, "p (pref)", border=1)
        pdf.ln()
        for crit in dados['tipos'].keys():
            crit_ascii = ascii_only(crit)
            tipo = dados['tipos'][crit]
            q_val = dados['q_limites'][crit]
            p_val = dados['p_limites'][crit]
            pdf.cell(40, 5, crit_ascii, border=1)
            pdf.cell(20, 5, tipo, border=1)
            pdf.cell(30, 5, f"{q_val:.2f}", border=1)
            pdf.cell(30, 5, f"{p_val:.2f}", border=1)
            pdf.ln()
        pdf.ln(10)

    adicionar_tabela_texto("Dados de Entrada:", dados['entrada'])
    adicionar_tabela_texto("Matriz Normalizada (CRITIC):", dados['norm'])
    if not dados['correl'].empty:
        adicionar_tabela_texto("Matriz de Correlação (CRITIC):", dados['correl'])
    adicionar_tabela_texto("Matriz de Preferência Agregada:", dados['pref_matrix'])

    pdf.set_font("Arial", 'B', 10)
    pdf.cell(200, 10, txt="Pesos dos Critérios:", ln=True)
    pdf.set_font("Arial", size=10)
    for crit, peso in dados['pesos'].items():
        crit_ascii = ascii_only(crit)
        pdf.cell(200, 5, txt=f"{crit_ascii}: {peso:.4f}", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 10)
    pdf.cell(200, 10, txt="Fluxos PROMETHEE:", ln=True)
    pdf.set_font("Arial", size=10)
    for alt in dados['phi_liquido'].index:
        alt_ascii = ascii_only(alt)
        pdf.cell(200, 5, txt=f"{alt_ascii}: Fluxo+ = {dados['phi_mais'][alt]:.4f}, Fluxo- = {dados['phi_menos'][alt]:.4f}, Fluxo Líquido = {dados['phi_liquido'][alt]:.4f}", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 10)
    pdf.cell(200, 10, txt="Ranking PROMETHEE II:", ln=True)
    pdf.set_font("Arial", size=10)
    for alt, pos in dados['ranking'].items():
        alt_ascii = ascii_only(alt)
        pdf.cell(200, 5, txt=f"{alt_ascii}: {pos}º", ln=True)
    pdf.ln(10)

    def adicionar_imagem(fig, titulo):
        titulo_ascii = ascii_only(titulo)
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name, format='png', bbox_inches='tight')
            tmp.close()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(200, 10, txt=titulo_ascii, ln=True, align='C')
            pdf.image(tmp.name, x=10, y=30, w=180)
            try:
                os.unlink(tmp.name)
            except PermissionError:
                pass

    adicionar_imagem(dados['fig_pesos'], "Gráfico de Pesos")
    adicionar_imagem(dados['fig_fluxos'], "Gráfico de Fluxos (barras)")
    adicionar_imagem(dados['fig_grafo'], "Grafo de Sobreclassificação (PROMETHEE I)")

    pdf.output(nome_arquivo)

def gerar_relatorio_docx(dados, nome_arquivo):
    doc = Document()
    doc.add_heading('Relatório MCDA - CRITIC + PROMETHEE', level=1)

    if 'q_limites' in dados and 'p_limites' in dados and 'tipos' in dados:
        doc.add_heading('Limiares dos Critérios (q e p)', level=2)
        tabela_lim = doc.add_table(rows=1, cols=4)
        tabela_lim.style = 'Light Grid Accent 1'
        hdr = tabela_lim.rows[0].cells
        hdr[0].text = 'Critério'
        hdr[1].text = 'Tipo'
        hdr[2].text = 'q (indiferença)'
        hdr[3].text = 'p (preferência)'
        for crit in dados['tipos'].keys():
            cells = tabela_lim.add_row().cells
            cells[0].text = str(crit)
            cells[1].text = dados['tipos'][crit]
            cells[2].text = f"{dados['q_limites'][crit]:.2f}"
            cells[3].text = f"{dados['p_limites'][crit]:.2f}"

    def adicionar_tabela_docx(titulo, df):
        doc.add_heading(titulo, level=2)
        tabela = doc.add_table(rows=1, cols=len(df.columns)+1)
        tabela.style = 'Light Grid Accent 1'
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Alternativa'
        for j, col in enumerate(df.columns):
            hdr_cells[j+1].text = str(col)
        for i, row in df.iterrows():
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(i)
            for j, val in enumerate(row):
                row_cells[j+1].text = f"{val:.4f}" if isinstance(val, (int, float)) else str(val)

    adicionar_tabela_docx("Dados de Entrada", dados['entrada'])
    adicionar_tabela_docx("Matriz Normalizada (CRITIC)", dados['norm'])
    if not dados['correl'].empty:
        adicionar_tabela_docx("Matriz de Correlação (CRITIC)", dados['correl'])
    adicionar_tabela_docx("Matriz de Preferência Agregada (Π)", dados['pref_matrix'])

    doc.add_heading('Pesos dos Critérios', level=2)
    for crit, peso in dados['pesos'].items():
        doc.add_paragraph(f"{crit}: {peso:.4f}")

    doc.add_heading('Fluxos PROMETHEE', level=2)
    tabela_fluxos = doc.add_table(rows=1, cols=4)
    tabela_fluxos.style = 'Light Grid Accent 1'
    hdr = tabela_fluxos.rows[0].cells
    hdr[0].text = 'Alternativa'
    hdr[1].text = 'Φ+'
    hdr[2].text = 'Φ-'
    hdr[3].text = 'Φ Líquido'
    for alt in dados['phi_liquido'].index:
        cells = tabela_fluxos.add_row().cells
        cells[0].text = str(alt)
        cells[1].text = f"{dados['phi_mais'][alt]:.4f}"
        cells[2].text = f"{dados['phi_menos'][alt]:.4f}"
        cells[3].text = f"{dados['phi_liquido'][alt]:.4f}"

    doc.add_heading('Ranking PROMETHEE II', level=2)
    for alt, pos in dados['ranking'].items():
        doc.add_paragraph(f"{alt}: {pos}º")

    def adicionar_imagem_docx(fig, titulo):
        doc.add_heading(titulo, level=2)
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name, format='png', bbox_inches='tight')
            tmp.seek(0)
            doc.add_picture(tmp.name, width=Inches(5))
            tmp.close()
            try:
                os.unlink(tmp.name)
            except PermissionError:
                pass

    adicionar_imagem_docx(dados['fig_pesos'], "Gráfico de Pesos")
    adicionar_imagem_docx(dados['fig_fluxos'], "Gráfico de Fluxos (barras)")
    adicionar_imagem_docx(dados['fig_grafo'], "Grafo de Sobreclassificação (PROMETHEE I)")

    doc.save(nome_arquivo)

# ------------------------------------------------------------
# FUNÇÃO PARA LIMPAR O ESTADO (NOVA ANÁLISE)
# ------------------------------------------------------------
def reset_app():
    keys_to_clear = ['dados_entrada', 'criterios_tipos', 'criterios_qs', 'criterios_ps',
                     'metodo_peso', 'pesos_manuais', 'resultados', 'params_confirmados']
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

# ------------------------------------------------------------
# INTERFACE STREAMLIT
# ------------------------------------------------------------

def main():
    st.set_page_config(page_title="MCDA - CRITIC + PROMETHEE", layout="wide")
    st.title("📊 Análise Multicritério (CRITIC + PROMETHEE I/II)")
    st.markdown("""
    Implementação rigorosa fundamentada na Escola Francesa de tomada de decisão.
    """)

    # Botão de Nova Análise no topo
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("🔄 Nova Análise", use_container_width=True):
            reset_app()
            st.rerun()

    # Inicialização de Variáveis de Estado
    if 'dados_entrada' not in st.session_state:
        st.session_state.dados_entrada = None
    if 'criterios_tipos' not in st.session_state:
        st.session_state.criterios_tipos = {}
    if 'criterios_qs' not in st.session_state:
        st.session_state.criterios_qs = {}
    if 'criterios_ps' not in st.session_state:
        st.session_state.criterios_ps = {}
    if 'metodo_peso' not in st.session_state:
        st.session_state.metodo_peso = "Calcular automaticamente (CRITIC)"
    if 'pesos_manuais' not in st.session_state:
        st.session_state.pesos_manuais = None
    if 'resultados' not in st.session_state:
        st.session_state.resultados = {}
    if 'params_confirmados' not in st.session_state:
        st.session_state.params_confirmados = False

    # 1. ENTRADA DE DADOS
    st.header("1. Entrada de Dados")
    opcao = st.radio("Escolha a forma de entrada:", ["Upload de arquivo (Excel/CSV)", "Inserir manualmente"])

    if opcao == "Upload de arquivo (Excel/CSV)":
        arquivo = st.file_uploader("Faça upload da planilha", type=['xlsx', 'csv'])
        if arquivo is not None:
            try:
                if arquivo.name.endswith('.csv'):
                    df = pd.read_csv(arquivo, index_col=0, encoding='utf-8')
                else:
                    df = pd.read_excel(arquivo, index_col=0)
                st.session_state.dados_entrada = df
                st.success("Arquivo carregado com sucesso!")
                st.dataframe(df)
            except Exception as e:
                st.error(f"Erro ao ler arquivo: {e}")
    else:
        st.subheader("Defina as dimensões da Matriz de Decisão")
        cols_dim = st.columns(2)
        num_alt = cols_dim[0].number_input("Número de Alternativas", min_value=2, value=3, step=1)
        num_crit = cols_dim[1].number_input("Número de Critérios", min_value=1, value=3, step=1)

        if num_alt and num_crit:
            nomes_alt = [st.sidebar.text_input(f"Alt {i+1}", value=f"A{i+1}") for i in range(num_alt)]
            nomes_crit = [st.sidebar.text_input(f"Crit {j+1}", value=f"C{j+1}") for j in range(num_crit)]

            dados = []
            st.write("**Preencha as avaliações:**")
            for i, alt in enumerate(nomes_alt):
                cols_val = st.columns(num_crit)
                linha = [cols_val[j].number_input(f"{alt} - {crit}", value=0.0, step=0.1, key=f"v_{i}_{j}") for j, crit in enumerate(nomes_crit)]
                dados.append(linha)

            if st.button("Carregar matriz manual"):
                st.session_state.dados_entrada = pd.DataFrame(dados, index=nomes_alt, columns=nomes_crit)
                st.success("Matriz carregada!")
                st.dataframe(st.session_state.dados_entrada)

    # 2. DEFINIÇÃO DE PARÂMETROS DOS CRITÉRIOS
    if st.session_state.dados_entrada is not None:
        st.header("2. Parâmetros dos Critérios (PROMETHEE)")
        st.markdown("""
        Defina o tipo (Max/Min) e os limiares para a **Função de Preferência Tipo V**:
        * **q (Indiferença):** Diferenças até este valor são ignoradas.
        * **p (Preferência):** Diferenças a partir deste valor geram preferência total (1.0).
        *(Deixe q=0 e p=0 para usar o Critério Usual - Tipo I)*
        """)

        df = st.session_state.dados_entrada
        criterios = df.columns.tolist()

        for crit in criterios:
            if crit not in st.session_state.criterios_tipos:
                st.session_state.criterios_tipos[crit] = "max"
            if crit not in st.session_state.criterios_qs:
                st.session_state.criterios_qs[crit] = 0.0
            if crit not in st.session_state.criterios_ps:
                st.session_state.criterios_ps[crit] = 0.0

        valid_limiares = True
        for crit in criterios:
            st.write(f"**Critério: {crit}**")
            cols_params = st.columns(3)

            tipo = cols_params[0].selectbox("Objetivo", ["max", "min"], 
                                             index=0 if st.session_state.criterios_tipos[crit]=="max" else 1,
                                             key=f"t_{crit}")
            q = cols_params[1].number_input("Limiar q (Indiferença)", min_value=0.0, 
                                             value=st.session_state.criterios_qs[crit], 
                                             step=0.01, key=f"q_{crit}")
            p = cols_params[2].number_input("Limiar p (Preferência Forte)", min_value=0.0, 
                                             value=st.session_state.criterios_ps[crit], 
                                             step=0.01, key=f"p_{crit}")

            st.session_state.criterios_tipos[crit] = tipo
            st.session_state.criterios_qs[crit] = q
            st.session_state.criterios_ps[crit] = p

            if p < q:
                st.error(f"O limiar 'p' deve ser maior ou igual a 'q' no critério {crit}.")
                valid_limiares = False

        col_conf1, col_conf2, col_conf3 = st.columns([1, 2, 1])
        with col_conf2:
            if st.button("✅ Confirmar Parâmetros e Avançar", disabled=not valid_limiares, use_container_width=True):
                st.session_state.params_confirmados = True
                st.success("Parâmetros confirmados! Agora vá para a seção 3.")

        st.subheader("Método de Definição dos Pesos")
        metodo_peso = st.radio("Escolha como definir os pesos dos critérios:",
                               ["Calcular automaticamente (CRITIC)", "Inserir manualmente"],
                               index=0 if st.session_state.metodo_peso == "Calcular automaticamente (CRITIC)" else 1)
        st.session_state.metodo_peso = metodo_peso

        if metodo_peso == "Inserir manualmente":
            st.markdown("Digite os pesos (valores entre 0 e 1). A soma será normalizada automaticamente.")
            pesos_manuais = {}
            total = 0.0
            cols_peso = st.columns(len(criterios))
            for i, crit in enumerate(criterios):
                with cols_peso[i]:
                    default = 1.0 / len(criterios)
                    key = f"peso_manual_{crit}"
                    if key not in st.session_state:
                        st.session_state[key] = default
                    peso = st.number_input(f"Peso {crit}", min_value=0.0, max_value=1.0, 
                                            value=st.session_state[key], step=0.05, key=key)
                    pesos_manuais[crit] = peso
                    total += peso
            st.write(f"Soma atual: {total:.4f}")
            if abs(total - 1.0) > 0.01:
                st.warning("A soma dos pesos não é igual a 1. Eles serão normalizados automaticamente.")
            if st.button("Confirmar pesos manuais"):
                if total > 0:
                    st.session_state.pesos_manuais = {k: v/total for k, v in pesos_manuais.items()}
                else:
                    st.error("A soma dos pesos deve ser maior que zero.")
                    st.session_state.pesos_manuais = None
                st.success("Pesos manuais definidos e normalizados!")

    # 3. EXECUÇÃO
    if st.session_state.dados_entrada is not None:
        st.header("3. Execução e Resultados")

        pode_executar = st.session_state.params_confirmados and valid_limiares
        if st.session_state.metodo_peso == "Inserir manualmente" and st.session_state.pesos_manuais is None:
            pode_executar = False
            st.warning("Defina os pesos manuais na seção 2 antes de executar.")

        if st.button("▶️ Rodar CRITIC + PROMETHEE", disabled=not pode_executar):
            df = st.session_state.dados_entrada
            tipos = st.session_state.criterios_tipos
            qs = st.session_state.criterios_qs
            ps = st.session_state.criterios_ps
            metodo_peso = st.session_state.metodo_peso

            with st.spinner("Realizando cálculos..."):
                df_norm = normalizar_critic(df, tipos)

                if metodo_peso == "Calcular automaticamente (CRITIC)":
                    pesos, desvios, correl, info_c = metodo_critic(df_norm)
                    st.subheader("Fase I: Método CRITIC (Pesos)")
                    col1, col2 = st.columns(2)
                    col1.write("**Índice Cj (Quantidade de Informação):**")
                    col1.dataframe(info_c.to_frame(name="Cj"))
                    col2.write("**Pesos Finais dos Atributos (Wj):**")
                    col2.dataframe(pd.Series(pesos, name="Peso Wj"))
                    st.pyplot(gerar_grafico_pesos(pesos))
                else:
                    pesos = st.session_state.pesos_manuais
                    correl = pd.DataFrame()
                    st.subheader("Fase I: Pesos Inseridos Manualmente")
                    st.write("**Pesos (normalizados):**")
                    st.json(pesos)
                    st.pyplot(gerar_grafico_pesos(pesos))

                st.subheader("Fase II: Método PROMETHEE I e II")
                pref_matrix, phi_mais, phi_menos, phi_liquido, ranking = promethee(df, pesos, tipos, qs, ps)

                st.write("**Matriz de Índice Global de Preferência (Π):**")
                st.dataframe(pref_matrix)

                st.write("**Fluxos de Superação e Ranking Final (PROMETHEE II):**")
                res_df = pd.DataFrame({
                    'Φ+ (Positivo)': phi_mais,
                    'Φ- (Negativo)': phi_menos,
                    'Φ (Líquido)': phi_liquido,
                    'Posição Ranking': ranking
                })
                st.dataframe(res_df.sort_values(by='Posição Ranking'))

                col_graf1, col_graf2 = st.columns(2)
                with col_graf1:
                    st.pyplot(gerar_grafico_fluxos(phi_mais, phi_menos, phi_liquido))
                with col_graf2:
                    st.pyplot(gerar_grafo_sobreclassificacao(phi_mais, phi_menos))

                st.session_state.resultados = {
                    'entrada': df,
                    'norm': df_norm,
                    'correl': correl,
                    'pref_matrix': pref_matrix,
                    'pesos': pesos,
                    'phi_mais': phi_mais,
                    'phi_menos': phi_menos,
                    'phi_liquido': phi_liquido,
                    'ranking': ranking,
                    'fig_pesos': gerar_grafico_pesos(pesos),
                    'fig_fluxos': gerar_grafico_fluxos(phi_mais, phi_menos, phi_liquido),
                    'fig_grafo': gerar_grafo_sobreclassificacao(phi_mais, phi_menos),
                    'q_limites': qs,
                    'p_limites': ps,
                    'tipos': tipos
                }

                st.success("Análise concluída!")

    # 4. RELATÓRIO
    if st.session_state.resultados:
        st.header("4. Gerar Relatório")
        formato = st.selectbox("Formato do relatório", ["PDF", "DOCX"])
        if st.button("📄 Gerar e baixar relatório"):
            resultados = st.session_state.resultados
            with st.spinner("Gerando relatório..."):
                dados_relatorio = {
                    'entrada': resultados['entrada'],
                    'norm': resultados['norm'],
                    'correl': resultados['correl'],
                    'pref_matrix': resultados['pref_matrix'],
                    'pesos': resultados['pesos'],
                    'phi_mais': resultados['phi_mais'],
                    'phi_menos': resultados['phi_menos'],
                    'phi_liquido': resultados['phi_liquido'],
                    'ranking': resultados['ranking'],
                    'fig_pesos': resultados['fig_pesos'],
                    'fig_fluxos': resultados['fig_fluxos'],
                    'fig_grafo': resultados['fig_grafo'],
                    'q_limites': resultados['q_limites'],
                    'p_limites': resultados['p_limites'],
                    'tipos': resultados['tipos']
                }
                if formato == "PDF":
                    arquivo_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                    gerar_relatorio_pdf(dados_relatorio, arquivo_temp.name)
                    with open(arquivo_temp.name, 'rb') as f:
                        st.download_button("📥 Baixar PDF", f, file_name="relatorio_mcda.pdf")
                    os.unlink(arquivo_temp.name)
                else:
                    arquivo_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                    gerar_relatorio_docx(dados_relatorio, arquivo_temp.name)
                    with open(arquivo_temp.name, 'rb') as f:
                        st.download_button("📥 Baixar DOCX", f, file_name="relatorio_mcda.docx")
                    os.unlink(arquivo_temp.name)

        # Botão Nova Análise no final
        st.divider()
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🔄 Nova Análise (Final)", use_container_width=True):
                reset_app()
                st.rerun()

if __name__ == "__main__":
    main()
